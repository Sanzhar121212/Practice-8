from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from jinja2 import Environment, FileSystemLoader, select_autoescape



# weasyprint и python-docx импортируются там, где реально нужны,
# чтобы тест "босс-файт" не тормозил.


TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
PARCHMENTS_DIR = Path(__file__).resolve().parent.parent / "parchments"
PARCHMENTS_DIR.mkdir(exist_ok=True)


@dataclass
class RenderResult:
    html: str
    qr_path: Optional[Path]


class TemplateEngine:
    def __init__(self, templates_dir: Path = TEMPLATES_DIR) -> None:
        self.env = Environment(
            loader=FileSystemLoader(str(templates_dir)),
            autoescape=select_autoescape(["html", "xml"]),
        )

    def _generate_qr(self, quest_id: int) -> Optional[Path]:

        url = f"https://guild.example.com/quests/{quest_id}"

        qr_dir = PARCHMENTS_DIR / "qr"
        qr_dir.mkdir(exist_ok=True)
        qr_path = qr_dir / f"quest_{quest_id}.png"

        return qr_path

    def render(self, quest: Dict[str, Any], template_name: str) -> RenderResult:
        template = self.env.get_template(template_name)
        qr_path = self._generate_qr(quest["id"])
        html = template.render(
            quest=quest,
            now=datetime.now(),
            qr_code_path=str(qr_path) if qr_path else None,
        )
        return RenderResult(html=html, qr_path=qr_path)

    def export_pdf(self, render: RenderResult, output_path: Path) -> None:
        from weasyprint import HTML  # локальный импорт

        HTML(string=render.html).write_pdf(str(output_path))

    def export_docx(self, quest: Dict[str, Any], output_path: Path) -> None:
        from docx import Document  # локальный импорт

        doc = Document()
        doc.add_heading(f"Квест: {quest['title']}", level=1)
        doc.add_paragraph(f"ID: {quest['id']}")
        doc.add_paragraph(f"Сложность: {quest['difficulty']}")
        doc.add_paragraph(f"Награда: {quest['reward']} золотых")
        doc.add_paragraph("Описание:")
        doc.add_paragraph(quest["description"])
        doc.add_paragraph(f"Дедлайн: {quest['deadline']}")
        doc.add_paragraph(f"Создано: {quest['created_at']}")
        doc.save(str(output_path))

    @staticmethod
    def default_output_path(quest_id: int, ext: str) -> Path:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        return PARCHMENTS_DIR / f"{quest_id}_{ts}.{ext}"


class BatchExporter:
    """Для теста 'босс-файт': просто быстро генерируем HTML 100 раз."""

    @staticmethod
    def generate_100_quests() -> List[str]:
        engine = TemplateEngine()
        html_results: List[str] = []
        fake_quest_template: Dict[str, Any] = {
            "id": 1,
            "title": "Test Quest",
            "difficulty": "Средний",
            "reward": 100,
            "description": "Lorem ipsum " * 20,
            "deadline": "2025-12-31 23:59",
            "created_at": "2025-01-01 00:00",
        }
        for i in range(100):
            quest = fake_quest_template.copy()
            quest["id"] = i + 1
            result = engine.render(quest, "guild_contract.html")
            html_results.append(result.html)
        return html_results
